"""
export.py — Génération des fichiers de sortie : TXT, DOCX, SRT.

Gère deux modes :
  - Simple     : texte brut (mode Cohere sans diarisation)
  - Diarisation : texte avec interlocuteurs et horodatages
    Format :  [HH:MM:SS] 🎙️ Intervenant N : texte...
"""

import os
import uuid
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from transcriber import TranscriptionResult


# ─── Helpers de formatage ──────────────────────────────────────────────────────

def _format_srt_time(seconds: float) -> str:
    """Convertit des secondes en format SRT : HH:MM:SS,mmm"""
    seconds = max(0.0, seconds)
    hours   = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs    = int(seconds % 60)
    millis  = int(round((seconds - int(seconds)) * 1000))
    if millis >= 1000:
        millis = 999
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def _format_timestamp(seconds: float) -> str:
    """Convertit des secondes en HH:MM:SS pour l'affichage."""
    seconds = max(0.0, seconds)
    hours   = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs    = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _format_duration(seconds: float) -> str:
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"{m}:{s:02d}"


def _has_speakers(result: "TranscriptionResult") -> bool:
    """Retourne True si la transcription contient des informations sur les interlocuteurs."""
    return any(seg.speaker for seg in result.segments)


# ─── Service d'export ─────────────────────────────────────────────────────────

class ExportService:

    # ── TXT ──────────────────────────────────────────────────────────────────

    def to_txt(self, result: "TranscriptionResult") -> str:
        """
        Génère le texte brut.
        Avec diarisation : [HH:MM:SS] 🎙️ Intervenant N : texte
        Sans diarisation : texte simple
        """
        if not _has_speakers(result):
            return result.text

        lines = []
        for seg in result.segments:
            ts   = _format_timestamp(seg.start)
            name = seg.speaker or "Inconnu"
            lines.append(f"[{ts}] 🎙️ {name} : {seg.text.strip()}")

        return "\n".join(lines)

    # ── SRT ──────────────────────────────────────────────────────────────────

    def to_srt(self, result: "TranscriptionResult") -> str:
        """
        Génère le fichier de sous-titres SRT.
        Avec diarisation : inclut le nom de l'interlocuteur dans le texte.
        """
        if not result.segments:
            return (
                f"1\n"
                f"{_format_srt_time(0.0)} --> {_format_srt_time(result.duration)}\n"
                f"{result.text.strip()}\n\n"
            )

        blocks = []
        idx = 1
        for seg in result.segments:
            text = seg.text.strip()
            if not text:
                continue

            # Préfixe interlocuteur si disponible
            if seg.speaker:
                text = f"[{seg.speaker}] {text}"

            block = (
                f"{idx}\n"
                f"{_format_srt_time(seg.start)} --> {_format_srt_time(seg.end)}\n"
                f"{text}\n"
            )
            blocks.append(block)
            idx += 1

        return "\n".join(blocks) + "\n"

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def to_docx(self, result: "TranscriptionResult", output_path: str) -> str:
        """
        Génère un document Word (.docx).
        Avec diarisation : chaque interlocuteur est mis en valeur avec sa couleur.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Titre
        title = doc.add_heading("Transcription", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Métadonnées
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(
            f"Langue : {result.language.upper()}  |  "
            f"Durée : {_format_duration(result.duration)}"
        )
        meta_run.font.size  = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()  # espace

        if _has_speakers(result):
            # Couleurs par interlocuteur (cycle)
            SPEAKER_COLORS = [
                RGBColor(0x7C, 0x3A, 0xED),  # violet
                RGBColor(0x0E, 0x7A, 0xBF),  # bleu
                RGBColor(0x0E, 0x8A, 0x4A),  # vert
                RGBColor(0xBF, 0x5E, 0x0E),  # orange
                RGBColor(0xBF, 0x0E, 0x5E),  # rose
                RGBColor(0x5E, 0x0E, 0xBF),  # indigo
            ]
            speaker_color_map: dict = {}
            color_idx = 0

            for seg in result.segments:
                text    = seg.text.strip()
                speaker = seg.speaker or "Inconnu"
                ts      = _format_timestamp(seg.start)

                if not text:
                    continue

                # Assigner une couleur unique à chaque interlocuteur
                if speaker not in speaker_color_map:
                    speaker_color_map[speaker] = SPEAKER_COLORS[color_idx % len(SPEAKER_COLORS)]
                    color_idx += 1

                color = speaker_color_map[speaker]

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)

                # Horodatage
                ts_run = p.add_run(f"[{ts}] ")
                ts_run.font.size  = Pt(9)
                ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                # Nom de l'interlocuteur
                speaker_run = p.add_run(f"🎙️ {speaker} : ")
                speaker_run.bold = True
                speaker_run.font.color.rgb = color
                speaker_run.font.size = Pt(10.5)

                # Texte
                text_run = p.add_run(text)
                text_run.font.size = Pt(10.5)

        else:
            # Mode simple : texte brut
            for line in result.text.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.paragraph_format.space_after = Pt(6)
                else:
                    doc.add_paragraph()

        doc.save(output_path)
        return output_path

    # ── Réponse JSON ─────────────────────────────────────────────────────────

    def generate_export_response(
        self,
        result: "TranscriptionResult",
        temp_dir: str,
    ) -> dict:
        """Génère les contenus TXT et SRT à embarquer dans la réponse JSON."""
        return {
            "txt": self.to_txt(result),
            "srt": self.to_srt(result),
        }
